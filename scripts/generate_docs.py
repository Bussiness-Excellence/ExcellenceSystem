import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

def build_word_doc():
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("PulpoPlus CRM Data Pipeline")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87) # Dark Blue

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("End-to-End Technical Roadmap & Pipeline Flow Milestones")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(80, 80, 80)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 44, 87)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(41, 128, 185)
        return p

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = "Calibri"
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(30, 30, 30)
        run_t = p.add_run(text)
        run_t.font.name = "Calibri"
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(50, 50, 50)
        return p

    # Overview
    add_heading_1("1. Pipeline Architecture Overview")
    add_body("This document details the end-to-end data processing workflow for PulpoPlus CRM data. The pipeline automates the transformation of raw HTML export files into structured analytical Excel workbooks and seamlessly synchronizes them with Supabase cloud database tables.")

    # Milestones
    add_heading_1("2. Milestone-by-Milestone Pipeline Flow")

    # Milestone 1
    add_heading_2("Milestone 1: Data Ingestion & HTML Extraction")
    add_body(" PulpoPlus CRM exports raw web data saved with `.xls` extension (HTML format).", "Source Input: ")
    add_body(" `pulpoplus_extract_visits.py`", "Executing Script: ")
    add_body(" The script parses all HTML tables across 5 core sections:", "Process Actions: ")
    doc.add_paragraph("   • Visits (Clinic, Polyclinic, Hospital, AM Center calls)\n   • Pharmacies Visits (Stock, order, and pharmacy visit logs)\n   • Office Work (Administrative and field prep work)\n   • Activities (Coaching, meetings, and team events)\n   • Events (Conferences, seminars, and medical events)")
    add_body(" Extracted raw records and cleaned individual visit entries.", "Output Artifact: ")

    # Milestone 2
    add_heading_2("Milestone 2: Data Deduplication & Business Intelligence Calculations")
    add_body(" Raw extracted records from Milestone 1.", "Source Input: ")
    add_body(" `pulpoplus_rebuild_summary.py` (imported automatically by `pulpoplus_extract_visits.py`)", "Executing Script: ")
    add_body(" Business logic algorithms calculate key operational metrics:", "Process Actions: ")
    doc.add_paragraph("   • Deduplication: Removes duplicated joint/coaching visit rows.\n   • Hierarchy & Territory Alignment: Maps employees to team structure.\n   • Coaching Days Calculation: Matches AM/PM manager-rep joint visits.\n   • Shift & Coverage Analytics: Calculates AM/PM field shift durations, start times, unique doctor counts, and specialty doctor coverage.\n   • Product Analytics: Computes product calls per specialty.")
    add_body(" Cleaned summary workbook (e.g. `July 2026 - Eagles 1.xlsx`) with 5 sheets: Summary, Raw Data, Coaching Days, Specialty x Class, Product Calls per spec.", "Output Artifact: ")

    # Milestone 3
    add_heading_2("Milestone 3: Cloud Synchronization & Supabase Storage")
    add_body(" Calculated Excel workbook (`.xlsx`) + Team Hierarchy file.", "Source Input: ")
    add_body(" `pulpoplus_upload_to_supabase.py` / `pulpoplus_auto_upload.py`", "Executing Script: ")
    add_body(" Rest API calls batch-insert structured data into Supabase:", "Process Actions: ")
    doc.add_paragraph("   • Hierarchy Sync: Updates `teams` and `hierarchy` tables.\n   • Batch Replace: Clears prior upload batch to avoid duplication.\n   • Table Insert: Populates `visits`, `coaching_days`, `summaries`, `specialty_classification`, and `product_calls` tables.")
    add_body(" Live cloud data ready for dashboard visualization.", "Output Artifact: ")

    # Milestone 4
    add_heading_2("Milestone 4: Web Application & Executive Dashboard Presentation")
    add_body(" Supabase Database tables.", "Source Input: ")
    add_body(" React Dashboard (`excellence-crm`) + Vercel API Route (`upload_visits.py`)", "Executing Application: ")
    add_body(" Web application delivers analytics UI and direct drag-and-drop web uploader at `/admin/upload`.", "Process Actions: ")
    add_body(" Interactive charts, KPI metrics, rep performance tables, and automated upload UI.", "Output Artifact: ")

    # Summary Table
    add_heading_1("3. Pipeline Roadmap Summary Table")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    headers = ["Milestone Stage", "Input Data", "Core Script / Logic", "Output Target"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="102C57"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    data_rows = [
        ("M1: HTML Extraction", "Raw HTML Export (.xls)", "pulpoplus_extract_visits.py", "Parsed Raw Records"),
        ("M2: Metric Calculation", "Parsed Records", "pulpoplus_rebuild_summary.py", "Summary .xlsx Workbook"),
        ("M3: Cloud Sync", "Summary Workbook", "pulpoplus_upload_to_supabase.py", "Supabase DB Tables"),
        ("M4: UI Presentation", "Supabase Tables", "React App / Dashboard.js", "Executive Dashboard"),
    ]

    for stage, inp, script, out in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = inp
        row_cells[2].text = script
        row_cells[3].text = out
        for cell in row_cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
            cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(40, 40, 40)

    doc.save("Pipeline_Flow_Roadmap_and_Milestones.docx")
    print("Word document created successfully: Pipeline_Flow_Roadmap_and_Milestones.docx")

def build_excel_roadmap():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Pipeline Roadmap"

    # Color Palette
    navy_fill = PatternFill(start_color="102C57", end_color="102C57", fill_type="solid")
    blue_fill = PatternFill(start_color="2980B9", end_color="2980B9", fill_type="solid")
    light_fill = PatternFill(start_color="F4F6F9", end_color="F4F6F9", fill_type="solid")
    white_bold = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    title_font = Font(name="Calibri", size=16, bold=True, color="102C57")
    section_font = Font(name="Calibri", size=12, bold=True, color="2980B9")
    regular_font = Font(name="Calibri", size=10, color="333333")

    thin_border = Border(
        left=Side(style='thin', color='D0D7DE'),
        right=Side(style='thin', color='D0D7DE'),
        top=Side(style='thin', color='D0D7DE'),
        bottom=Side(style='thin', color='D0D7DE')
    )

    ws['A1'] = "PulpoPlus CRM Data Pipeline - End-to-End Roadmap"
    ws['A1'].font = title_font
    ws.merge_cells('A1:F1')

    ws['A2'] = "Milestone-by-Milestone Technical Flow & Operations"
    ws['A2'].font = section_font
    ws.merge_cells('A2:F2')

    headers = ["Milestone #", "Stage Name", "Input Artifact", "Primary Python / System Script", "Core Operations & Calculations", "Output Target Artifact"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col_num)
        cell.value = header
        cell.fill = navy_fill
        cell.font = white_bold
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    rows = [
        ("Milestone 1", "HTML Report Extraction", "Raw HTML Export (.xls)", "pulpoplus_extract_visits.py", "Parses 5 HTML report sections: Visits, Pharmacy Visits, Office Work, Activities, Events. Decodes encoded HTML tables.", "Parsed Records Array"),
        ("Milestone 2", "Deduplication & Metrics Calculation", "Parsed Records Array + Team Structure.xlsx", "pulpoplus_rebuild_summary.py", "Deduplicates joint double-visits. Calculates AM/PM shift durations, field days, coaching days, doctor coverage & product call counts.", "Summary Workbook (.xlsx) with 5 sheets"),
        ("Milestone 3", "Cloud Database Synchronization", "Summary Workbook (.xlsx)", "pulpoplus_upload_to_supabase.py / pulpoplus_auto_upload.py", "Deletes prior batch data to avoid duplicates. Inserts rows into Supabase tables (visits, summaries, coaching_days, etc.).", "Supabase Database Tables"),
        ("Milestone 4", "Web Presentation & Direct Upload", "Supabase Tables / Raw File Upload", "React Uploader (src/pages/AdminUpload.js) + api/upload_visits.py", "Renders interactive analytics dashboard and provides web drag-and-drop file upload for non-technical users.", "Live Executive Dashboard"),
    ]

    for row_idx, data in enumerate(rows, 5):
        for col_idx, val in enumerate(data, 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = regular_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if row_idx % 2 == 1:
                cell.fill = light_fill

    # Set column widths
    widths = [15, 25, 25, 30, 45, 30]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    wb.save("Pipeline_Flow_Roadmap_and_Milestones.xlsx")
    print("Excel document created successfully: Pipeline_Flow_Roadmap_and_Milestones.xlsx")

if __name__ == "__main__":
    build_word_doc()
    build_excel_roadmap()
